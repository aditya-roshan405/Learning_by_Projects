"""
report_generator.py
Creates a simple Word (.docx) report summarizing the resume analysis
so the user can download and keep it.
"""

from docx import Document


def generate_report(
    candidate_skills: list[str],
    matched_role: str,
    match_score: float,
    missing_skills: list[str],
    roadmap: dict[str, str],
    output_path: str = "reports/resume_report.docx",
) -> str:
    """
    Build a Word document with the analysis results and save it.
    Returns the path where the report was saved.
    """
    doc = Document()

    doc.add_heading("AI Resume Analysis Report", level=1)

    doc.add_heading("Best Matching Role", level=2)
    doc.add_paragraph(f"{matched_role} - Match Score: {match_score}%")

    doc.add_heading("Skills Found in Resume", level=2)
    if candidate_skills:
        for skill in candidate_skills:
            doc.add_paragraph(skill, style="List Bullet")
    else:
        doc.add_paragraph("No known skills detected.")

    doc.add_heading("Missing Skills for this Role", level=2)
    if missing_skills:
        for skill in missing_skills:
            doc.add_paragraph(skill, style="List Bullet")
    else:
        doc.add_paragraph("Great! No major skill gaps found.")

    doc.add_heading("Suggested Learning Roadmap", level=2)
    for skill, tip in roadmap.items():
        doc.add_paragraph(f"{skill}: {tip}", style="List Bullet")

    doc.save(output_path)
    return output_path


# quick test
if __name__ == "__main__":
    path = generate_report(
        candidate_skills=["python", "sql"],
        matched_role="Data Analyst",
        match_score=78.5,
        missing_skills=["power bi"],
        roadmap={"power bi": "Microsoft's free Power BI learning path"},
    )
    print("Report saved at:", path)
